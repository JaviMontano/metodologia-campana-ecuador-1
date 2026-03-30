#!/usr/bin/env python3
"""
doc_converter.py — Multi-format document converter for MetodologIA
Converts between Markdown, DOCX, XLSX, PDF, PPTX, and HTML.

Usage:
    python doc_converter.py md2docx input.md output.docx
    python doc_converter.py md2pdf  input.md output.pdf
    python doc_converter.py md2xlsx input.md output.xlsx  (tables only)
    python doc_converter.py xlsx2md input.xlsx output.md
    python doc_converter.py docx2md input.docx output.md
    python doc_converter.py md2html input.md output.html

Dependencies: pip install -r requirements.txt
All imports are lazy — only loads what's needed for the requested conversion.
"""

import sys
import re
from pathlib import Path

# ── Brand Constants ─────────────────────────────────────────────────────────

BRAND = {
    'navy':    '#122562',
    'gold':    '#FFD700',
    'blue':    '#137DC5',
    'dark':    '#1F2833',
    'gray':    '#808080',
    'white':   '#FFFFFF',
    'light':   '#F8F9FC',
    'danger':  '#C0392B',
    'success': '#27AE60',
    'warning': '#E67E22',
    'font_heading': 'Arial',
    'font_body':    'Trebuchet MS',
}


def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple."""
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


# ── Markdown Parser ─────────────────────────────────────────────────────────

def parse_md_sections(md_text: str) -> list:
    """Parse markdown into sections with title, level, and content."""
    sections = []
    current = None
    for line in md_text.split('\n'):
        level = 0
        if line.startswith('### '):
            level, title = 3, line[4:]
        elif line.startswith('## '):
            level, title = 2, line[3:]
        elif line.startswith('# '):
            level, title = 1, line[2:]

        if level:
            current = {'level': level, 'title': title.strip(), 'content': []}
            sections.append(current)
        elif current is not None:
            current['content'].append(line)
    return sections


def parse_md_tables(md_text: str) -> list:
    """Extract all markdown tables as list of {headers, rows}."""
    tables = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            headers = [c.strip() for c in line.split('|')[1:-1]]
            # Skip separator
            if i + 1 < len(lines) and '---' in lines[i + 1]:
                i += 2
            else:
                i += 1
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                rows.append(cells)
                i += 1
            tables.append({'headers': headers, 'rows': rows})
        else:
            i += 1
    return tables


# ── MD → DOCX ───────────────────────────────────────────────────────────────

def md2docx(input_path: str, output_path: str):
    """Convert markdown to branded DOCX."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md = Path(input_path).read_text(encoding='utf-8')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = BRAND['font_body']
    font.size = Pt(11)
    font.color.rgb = RGBColor(*hex_to_rgb(BRAND['dark']))

    sections = parse_md_sections(md)
    for sec in sections:
        heading = doc.add_heading(sec['title'], level=min(sec['level'], 4))
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*hex_to_rgb(BRAND['navy']))

        for line in sec['content']:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+[.)]\s', stripped):
                text = re.sub(r'^\d+[.)]\s+', '', stripped)
                doc.add_paragraph(text, style='List Number')
            elif stripped.startswith('|'):
                continue  # Tables handled separately
            else:
                p = doc.add_paragraph()
                # Handle bold markers
                parts = re.split(r'(\*\*.+?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    # Add tables
    tables = parse_md_tables(md)
    for tbl in tables:
        if not tbl['headers']:
            continue
        table = doc.add_table(rows=1, cols=len(tbl['headers']))
        table.style = 'Table Grid'

        for i, header in enumerate(tbl['headers']):
            cell = table.rows[0].cells[i]
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for row_data in tbl['rows']:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = cell_text

    doc.save(output_path)
    print(f'DOCX saved: {output_path}')


# ── MD → PDF ────────────────────────────────────────────────────────────────

def md2pdf(input_path: str, output_path: str):
    """Convert markdown to branded PDF via reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RLTable, TableStyle

    md = Path(input_path).read_text(encoding='utf-8')
    sections = parse_md_sections(md)

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=25*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='BrandH1', fontName='Helvetica-Bold', fontSize=22,
                              textColor=HexColor(BRAND['navy']), spaceAfter=12))
    styles.add(ParagraphStyle(name='BrandH2', fontName='Helvetica-Bold', fontSize=16,
                              textColor=HexColor(BRAND['navy']), spaceBefore=18, spaceAfter=8))
    styles.add(ParagraphStyle(name='BrandH3', fontName='Helvetica-Bold', fontSize=13,
                              textColor=HexColor(BRAND['blue']), spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name='BrandBody', fontName='Helvetica', fontSize=10,
                              textColor=HexColor(BRAND['dark']), leading=14))

    story = []
    heading_styles = {1: 'BrandH1', 2: 'BrandH2', 3: 'BrandH3'}

    for sec in sections:
        style_name = heading_styles.get(sec['level'], 'BrandH3')
        story.append(Paragraph(sec['title'], styles[style_name]))

        for line in sec['content']:
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
                continue
            if stripped.startswith('|'):
                continue  # handled below
            text = stripped.replace('**', '<b>').replace('**', '</b>')
            if stripped.startswith('- ') or stripped.startswith('* '):
                text = '&bull; ' + text[2:]
            story.append(Paragraph(text, styles['BrandBody']))

    # Tables
    tables = parse_md_tables(md)
    for tbl in tables:
        if not tbl['headers']:
            continue
        data = [tbl['headers']] + tbl['rows']
        t = RLTable(data, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor(BRAND['navy'])),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor(BRAND['white'])),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CCCCCC')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor(BRAND['white']), HexColor(BRAND['light'])]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(Spacer(1, 12))
        story.append(t)

    doc.build(story)
    print(f'PDF saved: {output_path}')


# ── MD → XLSX (tables only) ─────────────────────────────────────────────────

def md2xlsx(input_path: str, output_path: str):
    """Extract markdown tables to branded XLSX."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    md = Path(input_path).read_text(encoding='utf-8')
    tables = parse_md_tables(md)

    wb = Workbook()
    ws = wb.active
    ws.title = 'Data'

    header_font = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill(start_color='122562', end_color='122562', fill_type='solid')
    body_font = Font(name='Trebuchet MS', size=10, color='1F2833')
    alt_fill = PatternFill(start_color='F8F9FC', end_color='F8F9FC', fill_type='solid')
    thin_border = Border(
        bottom=Side(style='thin', color='DDDDDD')
    )

    row_idx = 1
    for tbl_idx, tbl in enumerate(tables):
        if tbl_idx > 0:
            row_idx += 2  # gap between tables

        # Headers
        for col, header in enumerate(tbl['headers'], 1):
            cell = ws.cell(row=row_idx, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='left', vertical='center')
        row_idx += 1

        # Data
        for r, row_data in enumerate(tbl['rows']):
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.font = body_font
                cell.border = thin_border
                if r % 2 == 1:
                    cell.fill = alt_fill
            row_idx += 1

    # Auto-width columns
    for col_cells in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col_cells)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 50)

    wb.save(output_path)
    print(f'XLSX saved: {output_path}')


# ── XLSX → MD ────────────────────────────────────────────────────────────────

def xlsx2md(input_path: str, output_path: str):
    """Convert XLSX sheets to markdown tables."""
    from openpyxl import load_workbook

    wb = load_workbook(input_path, read_only=True)
    lines = [f'# {Path(input_path).stem}', '']

    for ws in wb.worksheets:
        lines.append(f'## {ws.title}')
        lines.append('')

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [str(c or '') for c in rows[0]]
        lines.append('| ' + ' | '.join(headers) + ' |')
        lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')

        for row in rows[1:]:
            cells = [str(c or '') for c in row]
            lines.append('| ' + ' | '.join(cells) + ' |')
        lines.append('')

    Path(output_path).write_text('\n'.join(lines), encoding='utf-8')
    print(f'MD saved: {output_path}')


# ── DOCX → MD ───────────────────────────────────────────────────────────────

def docx2md(input_path: str, output_path: str):
    """Convert DOCX to markdown."""
    from docx import Document

    doc = Document(input_path)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        if 'heading 1' in style:
            lines.append(f'# {text}')
        elif 'heading 2' in style:
            lines.append(f'## {text}')
        elif 'heading 3' in style:
            lines.append(f'### {text}')
        elif 'list bullet' in style:
            lines.append(f'- {text}')
        elif 'list number' in style:
            lines.append(f'1. {text}')
        else:
            # Reconstruct bold
            parts = []
            for run in para.runs:
                if run.bold:
                    parts.append(f'**{run.text}**')
                else:
                    parts.append(run.text)
            lines.append(''.join(parts))

    Path(output_path).write_text('\n'.join(lines), encoding='utf-8')
    print(f'MD saved: {output_path}')


# ── MD → HTML ───────────────────────────────────────────────────────────────

def md2html(input_path: str, output_path: str):
    """Convert markdown to standalone branded HTML."""
    try:
        import markdown
        md = Path(input_path).read_text(encoding='utf-8')
        body = markdown.markdown(md, extensions=['tables', 'fenced_code', 'toc'])
    except ImportError:
        # Fallback: basic conversion
        md = Path(input_path).read_text(encoding='utf-8')
        body = '<br>\n'.join(
            line.replace('# ', '<h1>').replace('## ', '<h2>').replace('### ', '<h3>')
            for line in md.split('\n')
        )

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{Path(input_path).stem}</title>
  <style>
    :root {{ --primary: {BRAND['navy']}; --accent: {BRAND['gold']}; --dark: {BRAND['dark']}; }}
    body {{ font-family: 'Trebuchet MS', Arial, sans-serif; color: var(--dark); max-width: 900px; margin: 0 auto; padding: 24px; line-height: 1.6; }}
    h1, h2, h3 {{ font-family: Arial, sans-serif; color: var(--primary); }}
    h2 {{ border-bottom: 3px solid var(--accent); padding-bottom: 4px; display: inline-block; }}
    table {{ width: 100%; border-collapse: collapse; margin: 16px 0; }}
    th {{ background: var(--primary); color: white; padding: 8px 12px; text-align: left; }}
    td {{ padding: 8px 12px; border-bottom: 1px solid rgba(0,0,0,0.08); }}
    tr:nth-child(even) {{ background: #F8F9FC; }}
    code {{ background: #F8F9FC; padding: 2px 6px; border-radius: 4px; font-size: 0.9em; }}
    pre {{ background: #F8F9FC; padding: 16px; border-radius: 8px; overflow-x: auto; }}
  </style>
</head>
<body>
{body}
</body>
</html>"""
    Path(output_path).write_text(html, encoding='utf-8')
    print(f'HTML saved: {output_path}')


# ── CLI ──────────────────────────────────────────────────────────────────────

CONVERTERS = {
    'md2docx': md2docx,
    'md2pdf':  md2pdf,
    'md2xlsx': md2xlsx,
    'md2html': md2html,
    'xlsx2md': xlsx2md,
    'docx2md': docx2md,
}

def main():
    if len(sys.argv) < 4:
        print('Usage: python doc_converter.py <command> <input> <output>')
        print(f'Commands: {", ".join(CONVERTERS.keys())}')
        sys.exit(1)

    cmd = sys.argv[1]
    if cmd not in CONVERTERS:
        print(f'Unknown command: {cmd}. Available: {", ".join(CONVERTERS.keys())}')
        sys.exit(1)

    CONVERTERS[cmd](sys.argv[2], sys.argv[3])


if __name__ == '__main__':
    main()
